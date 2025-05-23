from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import docx
from bs4 import BeautifulSoup, NavigableString, Tag
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run


@dataclass
class SimpleHtmlToDocx:
    """
    Convert restricted (simple, markdown-generated) HTML to DOCX format.
    """

    list_indent: float = 0.5  # inches
    max_indent: float = 5.5  # inches
    max_recursion_depth: int = 100  # Prevent stack overflow on deeply nested documents

    def convert_html_string(self, html: str) -> Document:
        """Convert HTML string to Document object."""
        soup = BeautifulSoup(html, "html.parser")
        doc = docx.Document()

        # Ensure 'Code' style exists
        self._ensure_style(doc, "Code")

        self._process_element(soup, doc)
        return doc

    def convert_html_file(
        self, input_path: Path | str, output_path: Path | str | None = None
    ) -> None:
        """Convert HTML file to DOCX file."""
        input_path = Path(input_path)
        try:
            html = input_path.read_text(encoding="utf-8")
            doc = self.convert_html_string(html)

            if output_path is None:
                output_path = input_path.with_suffix(".docx")

            doc.save(str(output_path))
        except Exception as e:
            raise RuntimeError(f"Error converting HTML file: {e}") from e

    def _process_element(
        self, element: Any, parent: Any, list_depth: int = 0, recursion_depth: int = 0
    ) -> None:
        """Recursively process HTML elements."""
        if recursion_depth > self.max_recursion_depth:
            return  # Prevent stack overflow

        if isinstance(element, NavigableString):
            # More refined whitespace handling - preserve single spaces but not unnecessary whitespace
            text = str(element)
            # Remove repeated whitespace but maintain basic structure
            text = re.sub(r"\s+", " ", text)
            if text and hasattr(parent, "add_run"):
                parent.add_run(text)
            return

        if not isinstance(element, Tag):
            return

        match element.name:
            case "h1" | "h2" | "h3" | "h4" | "h5" | "h6":
                level = int(element.name[1])
                p = parent.add_heading(level=level)
                self._process_children(element, p, list_depth, recursion_depth + 1)

            case "p":
                # If parent is already a paragraph, process children directly
                if isinstance(parent, Paragraph):
                    self._process_children(element, parent, list_depth, recursion_depth + 1)
                # Otherwise add a new paragraph to the parent
                elif hasattr(parent, "add_paragraph"):
                    p = parent.add_paragraph()
                    self._process_children(element, p, list_depth, recursion_depth + 1)
                # Fallback for unsupported parent
                else:
                    self._process_children(element, parent, list_depth, recursion_depth + 1)

            case "blockquote":
                p = parent.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(10)
                self._process_children(element, p, list_depth, recursion_depth + 1)

            case "pre":
                p = parent.add_paragraph()
                p.style = "Code"
                text = element.get_text().strip()
                p.add_run(text)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)

            case "code":
                if hasattr(parent, "add_run"):
                    run = parent.add_run(element.get_text().strip())
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                elif hasattr(parent, "add_paragraph"):
                    p = parent.add_paragraph()
                    p.style = "Code"
                    p.add_run(element.get_text().strip())
                else:
                    # We're in an unknown container that can't add runs or paragraphs
                    # Just convert to text as a fallback
                    text = element.get_text().strip()
                    if hasattr(parent, "text"):
                        parent.text = text

            case "ul" | "ol":
                for li in element.find_all("li", recursive=False):
                    if isinstance(li, Tag):
                        self._process_list_item(
                            li, parent, element.name, list_depth, recursion_depth + 1
                        )

            case "table":
                self._process_table(element, parent)

            case "hr":
                self._add_horizontal_rule(parent)

            case "br":
                if hasattr(parent, "add_run"):
                    parent.add_run().add_break()

            case "a" | "strong" | "b" | "em" | "i" | "code":
                if hasattr(parent, "add_run"):
                    # Preserve exact text without adding extra spaces
                    text = element.get_text()

                    # Apply appropriate formatting
                    run = parent.add_run(text)

                    if element.name in ["strong", "b"]:
                        run.font.bold = True
                    elif element.name in ["em", "i"]:
                        run.font.italic = True
                    elif element.name == "code":
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                    elif element.name == "a":
                        run.font.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)

                        # Handle href for links
                        href = element.get("href", "")
                        if isinstance(href, list):
                            href = " ".join(href)
                        href = str(href)
                        if href:
                            self._add_hyperlink(parent, run, href, text)

            case "img":
                # Support for image would go here if needed
                # Would require additional code to download images and add them to the document
                pass

            case _:
                # For any other tags, just process children
                self._process_children(element, parent, list_depth, recursion_depth + 1)

    def _process_children(
        self, element: Tag, parent: Any, list_depth: int = 0, recursion_depth: int = 0
    ) -> None:
        """Process all children of an element."""
        for child in element.children:
            # If the parent is a Run, pass the Run as parent to _process_element
            # This allows text and nested inline elements to be added to the same run
            if isinstance(parent, Run) and isinstance(child, NavigableString):
                parent.text += str(child)  # Append text to existing run
            elif isinstance(parent, Run) and isinstance(child, Tag):
                # For nested tags within a run, this gets more complex.
                # For simplicity, we'll assume here that inline tags don't nest deeply or complexly.
                # A more robust solution might need to break runs.
                self._process_element(child, parent, list_depth, recursion_depth + 1)
            else:
                self._process_element(child, parent, list_depth, recursion_depth + 1)

    def _process_list_item(
        self, li: Tag, doc: Document, list_type: str, depth: int, recursion_depth: int = 0
    ) -> None:
        """Process a list item with proper semantic list structure for nested lists."""
        # Choose list style based on type and depth
        if list_type == "ol":
            # Use the appropriate level of numbered list style
            if depth == 0:
                style_name = "List Number"
            elif depth == 1:
                style_name = "List Number 2"
            else:
                style_name = "List Number 3"  # Most Word templates support up to level 3
        else:  # ul
            # Use the appropriate level of bullet list style
            if depth == 0:
                style_name = "List Bullet"
            elif depth == 1:
                style_name = "List Bullet 2"
            else:
                style_name = "List Bullet 3"  # Most Word templates support up to level 3

        # Create the paragraph with the appropriate list style
        p = doc.add_paragraph(style=style_name)

        # Process direct content (excluding nested lists) into this paragraph
        direct_content = []
        nested_lists = []

        # Separate direct content from nested lists
        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                nested_lists.append(child)
            else:
                direct_content.append(child)

        # Process direct content
        for content in direct_content:
            self._process_element(content, p, depth, recursion_depth + 1)

        # Process nested lists (each gets its own paragraph with proper style)
        for nested_list in nested_lists:
            for nested_li in nested_list.find_all("li", recursive=False):
                if isinstance(nested_li, Tag):
                    self._process_list_item(
                        nested_li, doc, nested_list.name, depth + 1, recursion_depth + 1
                    )

    def _process_table(self, table: Tag, doc: Document) -> None:
        """Process a table element."""
        rows_elements = table.find_all("tr")
        if not rows_elements:
            return

        # Filter for actual Tag objects for rows
        rows_tags = [r for r in rows_elements if isinstance(r, Tag)]
        if not rows_tags:
            return

        # Calculate dimensions
        max_cols = max(len(row.find_all(["td", "th"])) for row in rows_tags)

        # Create table
        docx_table = doc.add_table(rows=len(rows_tags), cols=max_cols if max_cols > 0 else 1)
        docx_table.style = "Table Grid"  # Add borders

        # Fill cells
        for row_idx, row_element in enumerate(rows_tags):
            cell_elements = row_element.find_all(["td", "th"])
            for col_idx, cell_element in enumerate(cell_elements):
                if not isinstance(cell_element, Tag):
                    continue

                docx_cell = docx_table.cell(row_idx, col_idx)
                docx_cell.text = ""  # Clear default paragraph
                p = docx_cell.add_paragraph()

                # Bold for header cells
                if cell_element.name == "th":
                    run = p.add_run("")
                    run.font.bold = True

                # Handle alignment
                self._apply_cell_alignment(cell_element, p)

                # Process cell content
                self._process_children(cell_element, p)

    def _apply_cell_alignment(self, cell_element: Tag, paragraph: Paragraph) -> None:
        """Extract and apply alignment from cell attributes."""
        # Check both align attribute and style attribute for alignment info
        align_attr = cell_element.get("align", "") or ""
        style_attr = cell_element.get("style", "") or ""

        # Convert to string if it's a list
        if isinstance(align_attr, list):
            align_attr = " ".join(align_attr)
        if isinstance(style_attr, list):
            style_attr = " ".join(style_attr)

        # Extract text-align from style attribute if present
        text_align = ""
        if style_attr:
            align_match = re.search(r"text-align:\s*(\w+)", style_attr)
            if align_match:
                text_align = align_match.group(1).lower()

        # Determine alignment
        if "center" in align_attr or text_align == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "right" in align_attr or text_align == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # else: leave at default (left alignment)

    def _add_horizontal_rule(self, doc: Document) -> None:
        """Add a horizontal rule to the document."""
        p = doc.add_paragraph("* * *")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def _add_hyperlink(self, paragraph: Paragraph, run: Any, url: str, text: str) -> None:
        """Add hyperlink to the document."""
        # Create a new run with the same text and formatting
        new_run = paragraph.add_run(text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        if hasattr(run, "font") and hasattr(run.font, "color") and hasattr(run.font.color, "rgb"):
            new_run.font.color.rgb = run.font.color.rgb

        # Remove the original run
        run_parent = run._element.getparent()
        run_parent.remove(run._element)

        # Add hyperlink relationship
        r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create hyperlink and add the new run's element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        hyperlink.append(new_run._element)

        # Add hyperlink to paragraph
        paragraph._p.append(hyperlink)

    def _ensure_style(self, doc: Document, style_name: str) -> None:
        """Ensure a style exists in the document, creating it if necessary."""
        try:
            doc.styles[style_name]
        except KeyError:
            # Create code style if it doesn't exist
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


## Tests


_SAMPLE_HTML = """
<h1>Document Title</h1>
<p>This is a <strong>bold</strong> paragraph with <em>italic</em> text and a <a href="https://example.com">link</a>.</p>
<h2>Lists</h2>
<ul>
    <li>Unordered item 1</li>
    <li>Unordered item 2
        <ul>
            <li>Nested item A</li>
            <li>Nested item B</li>
        </ul>
    </li>
</ul>
<ol>
    <li>Ordered item 1</li>
    <li>Ordered item 2</li>
</ol>
<h2>Code</h2>
<pre><code>def hello_world():
    print("Hello, World!")
</code></pre>
<p>Inline <code>code</code> example</p>
<h2>Blockquote</h2>
<blockquote>
    <p>This is a blockquote with <strong>formatting</strong>.</p>
    <p>And a second paragraph.</p>
</blockquote>
<hr>
<h2>Table</h2>
<table>
    <tr>
        <th>Header 1</th>
        <th>Header 2</th>
        <th style="text-align:right">Right Aligned</th>
    </tr>
    <tr>
        <td>Row 1, Col 1</td>
        <td>Row 1, Col 2</td>
        <td align="right">1234</td>
    </tr>
    <tr>
        <td>Row 2, Col 1</td>
        <td>Row 2, Col 2</td>
        <td align="right">5678</td>
    </tr>
</table>
"""


# FIXME: Not roundtripping Markdown nested lists correctly.
# Fix code and then nested items below to be indented.
_EXPECTED_MD = r"""
# Document Title

This is a **bold** paragraph with *italic* text and a [link](https://example.com).

## Lists

* Unordered item 1
* Unordered item 2
* Nested item A
* Nested item B
1. Ordered item 1
2. Ordered item 2

## Code

def hello\_world():\
 print("Hello, World!")

Inline code example

## Blockquote

 This is a blockquote with **formatting**. And a second paragraph.

\* \* \*

## Table

| Header 1 | Header 2 | Right Aligned |
| --- | --- | --- |
| Row 1, Col 1 | Row 1, Col 2 | 1234 |
| Row 2, Col 1 | Row 2, Col 2 | 5678 |
"""


def test_html_to_docx_conversion():
    import os
    import tempfile

    from kash.kits.docs.doc_formats import docx_convert

    converter = SimpleHtmlToDocx()

    # Test string conversion
    doc = converter.convert_html_string(_SAMPLE_HTML)
    assert doc is not None

    # Verify document has content
    assert len(doc.paragraphs) > 0
    assert len(doc.tables) > 0

    # Test file conversion
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as temp_html:
        temp_html.write(_SAMPLE_HTML.encode("utf-8"))
        temp_html_path = temp_html.name

    temp_docx_path = Path(temp_html_path.replace(".html", ".docx"))
    try:
        converter.convert_html_file(temp_html_path, temp_docx_path)

        # Verify file exists and has content
        assert temp_docx_path.exists()
        assert temp_docx_path.stat().st_size > 0

        md = docx_convert.docx_to_md(temp_docx_path)

        print(md.markdown)

        assert md.markdown == _EXPECTED_MD.strip()

    finally:
        # Clean up temp files
        if os.path.exists(temp_html_path):
            os.unlink(temp_html_path)
        if os.path.exists(temp_docx_path):
            os.unlink(temp_docx_path)
